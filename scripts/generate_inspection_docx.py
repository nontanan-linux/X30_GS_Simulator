import os
import json
from docx import Document
from docx.shared import Inches

def generate_docx(mapping_json, images_dir, output_file, title="Inspection Points"):
    with open(mapping_json, 'r') as f:
        mapping = json.load(f)
        
    doc = Document()
    doc.add_heading(title, 0)
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Point Name'
    hdr_cells[1].text = 'Picture'
    
    for item in mapping:
        row_cells = table.add_row().cells
        row_cells[0].text = item['name']
        
        img_path = os.path.join(images_dir, item['file'])
        if os.path.exists(img_path):
            paragraph = row_cells[1].paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(img_path, width=Inches(3))
            
    doc.save(output_file)
    print(f"Generated {output_file}")

if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("--mapping", required=True)
    parser.add_argument("--images", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--title", default="Inspection Points")
    args = parser.parse_args()

    generate_docx(
        args.mapping,
        args.images,
        args.output,
        args.title
    )
